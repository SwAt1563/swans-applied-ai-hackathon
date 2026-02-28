import os
import json
from datetime import datetime, timedelta
from typing import Optional, Dict, Any, List
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()

from clio_sdk import clio
from pdf_parser import get_pdf_parser, AccidentDetails
from email_service import email_service
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from docx.shared import Pt, Inches  # <-- Add Inches here

# ==========================================
# Hackathon Multi-Tenant Mock
# For a real app, this would be extracted from a JWT or Session Cookie!
# ==========================================
DEMO_USER_ID = "demo_firm_account_1"

class ExtractedDataResponse(BaseModel):
    success: bool
    data: Optional[Dict[str, Any]] = None
    error: Optional[str] = None

class VerifiedData(BaseModel):
    matter_id: int
    date_of_accident: str
    accident_location: str
    defendant_name: str
    client_name: str
    client_vehicle_plate: str
    defendant_vehicle_plate: Optional[str] = None
    number_injured: int
    accident_description: str
    client_gender: str
    police_report_number: Optional[str] = None

class WorkflowResponse(BaseModel):
    success: bool
    matter_updated: bool = False
    custom_fields_set: bool = False
    calendar_entry_created: bool = False
    document_generated: bool = False
    email_sent: bool = False
    email_preview: Optional[Dict] = None
    errors: List[str] = []

@asynccontextmanager
async def lifespan(app: FastAPI):
    yield
    await clio.close_sdk()

app = FastAPI(title="Richards & Law - Police Report Automation", lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

@app.get("/oauth/login")
def oauth_login():
    # Pass the user ID into the URL generation so it is preserved in the OAuth state
    return RedirectResponse(url=clio.get_authorization_url(user_id=DEMO_USER_ID))

@app.get("/oauth/callback")
async def oauth_callback(code: str = Query(...), state: str = Query(DEMO_USER_ID)):
    try:
        # Save the tokens specifically for this user_id (passed back via 'state')
        await clio.exchange_code_for_tokens(code, user_id=state)
        return HTMLResponse(content="<h1>✓ Authorization Successful!</h1><p><a href='/ui'>Go to UI</a></p>")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"OAuth error: {str(e)}")

@app.get("/oauth/status")
def oauth_status():
    import time
    # Safely check the specific user's token file
    tokens = clio._read_tokens_from_file(DEMO_USER_ID)
    if not tokens or not tokens.get("access_token"):
        return {"authenticated": False}
    
    is_valid = time.time() < tokens.get("expires_at", 0)
    return {"authenticated": True, "token_valid": is_valid}

@app.post("/api/extract-pdf", response_model=ExtractedDataResponse)
async def extract_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are accepted")
    try:
        pdf_content = await file.read()
        parser = get_pdf_parser()
        accident_details = await parser.parse_police_report(pdf_content)
        data = accident_details.model_dump()
        data["statute_of_limitations_date"] = accident_details.statute_of_limitations_date
        return ExtractedDataResponse(success=True, data=data)
    except Exception as e:
        return ExtractedDataResponse(success=False, error=str(e))

@app.get("/api/matters")
async def get_matters():
    # Fetch matters using this user's specific access token
    return {"success": True, "data": await clio.get_matters(user_id=DEMO_USER_ID)}

@app.get("/api/document-templates")
async def get_document_templates():
    return {"success": True, "data": await clio.get_document_templates(user_id=DEMO_USER_ID)}

@app.post("/api/verify")
async def submit_verified_data(data: VerifiedData):
    try:
        custom_fields = await clio.get_custom_fields(user_id=DEMO_USER_ID, parent_type="Matter")
        field_map = {f["name"]: f["id"] for f in custom_fields}
        
        field_mappings = {
            "Date of Accident": data.date_of_accident,
            "Accident Location": data.accident_location,
            "Defendant Name": data.defendant_name,
            "Client Vehicle Plate": data.client_vehicle_plate,
            "Number Injured": str(data.number_injured),
            "Accident Description": data.accident_description,
            "Statute of Limitations": (datetime.strptime(data.date_of_accident, "%Y-%m-%d") + timedelta(days=8*365)).strftime("%Y-%m-%d"),
        }
        if data.defendant_vehicle_plate: field_mappings["Defendant Vehicle Plate"] = data.defendant_vehicle_plate
        if data.police_report_number: field_mappings["Police Report Number"] = data.police_report_number
        
        # Build {custom_field_id: value} mapping
        upsert_map = {}
        for field_name, value in field_mappings.items():
            field_id = field_map.get(field_name)
            if field_id:
                upsert_map[field_id] = value
                
        if upsert_map:
            await clio.upsert_matter_custom_fields(user_id=DEMO_USER_ID, matter_id=data.matter_id, field_id_value_map=upsert_map)
            
        return {"success": True, "results": {"custom_fields_updated": True}, "errors": []}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/workflow", response_model=WorkflowResponse)
async def run_full_workflow(matter_id: int = Form(...), template_id: int = Form(...), file: UploadFile = File(...)):
    errors = []
    response = WorkflowResponse(success=False)
    
    try:
        # Step 1: Extract PDF
        parser = get_pdf_parser()
        accident_details = await parser.parse_police_report(await file.read())
        
        # Step 2: Get Client
        matter = await clio.get_matter(user_id=DEMO_USER_ID, matter_id=matter_id, fields="id,client,responsible_attorney")
        client_info = matter.get("client", {})
        contact_id = client_info.get("id") if isinstance(client_info, dict) else None
        
        if contact_id:
            contact = await clio.get_contact(user_id=DEMO_USER_ID, contact_id=contact_id, fields="id,first_name,primary_email_address")
            email_field = contact.get("primary_email_address", "")
            client_email = email_field.get("address", "") if isinstance(email_field, dict) else email_field
            client_first_name = contact.get("first_name", accident_details.client_name.split()[0])
        else:
            client_email = ""
            client_first_name = accident_details.client_name.split()[0]
            errors.append("No client contact found on matter")
            
        # Step 3: Upsert Custom Fields
        sol_date = accident_details.statute_of_limitations_date
        
        # Calculate the Injury Clause paragraph dynamically!
        if accident_details.number_injured > 0:
            injury_clause = "Additionally, since the motor vehicle accident involved an injured person, Attorney will also investigate potential bodily injury claims and review relevant medical records to substantiate non-economic damages."
        else:
            injury_clause = "However, since the motor vehicle accident involved no reported injured people, the scope of this engagement is strictly limited to the recovery of property damage and loss of use."
        
        required_fields = {
            "Date of Accident": "date", 
            "Accident Location": "text_line", 
            "Defendant Name": "text_line", 
            "Client Vehicle Plate": "text_line", 
            "Defendant Vehicle Plate": "text_line",
            "Number Injured": "text_line", 
            "Accident Description": "text_area", 
            "Police Report Number": "text_line",
            "Statute of Limitations": "date",
            
            # --- NEW DYNAMIC FIELDS ---
            "Client Pronoun Subject": "text_line",       # he / she
            "Client Pronoun Possessive": "text_line",    # his / her
            "Injury Clause": "text_area"                 # The dynamic paragraph
        }
        
        field_map = await clio.ensure_custom_fields_exist(user_id=DEMO_USER_ID, required_fields=required_fields, parent_type="Matter")
        
        field_values = {
            "Date of Accident": accident_details.date_of_accident, 
            "Accident Location": accident_details.accident_location,
            "Defendant Name": accident_details.defendant_name, 
            "Client Vehicle Plate": accident_details.client_vehicle_plate,
            "Defendant Vehicle Plate": accident_details.defendant_vehicle_plate or "Unknown",
            "Number Injured": str(accident_details.number_injured), 
            "Accident Description": accident_details.accident_description,
            "Police Report Number": accident_details.police_report_number or "Unknown",
            "Statute of Limitations": sol_date,
            
            # --- POPULATE THE NEW DYNAMIC FIELDS ---
            "Client Pronoun Subject": accident_details.pronoun_he_she,
            "Client Pronoun Possessive": accident_details.pronoun_his_her,
            "Injury Clause": injury_clause
        }
        
        override_map = {field_map[k]: v for k, v in field_values.items() if k in field_map}
        try:
            await clio.upsert_matter_custom_fields(user_id=DEMO_USER_ID, matter_id=matter_id, field_id_value_map=override_map)
            response.custom_fields_set = True
            response.matter_updated = True
        except Exception as e:
            errors.append(f"Custom fields error: {str(e)}")
            
        # Step 4: Calendar Entry
        try:
            responsible_attorney = matter.get("responsible_attorney")
            attorney_user_id = responsible_attorney.get("id") if responsible_attorney else None
            
            calendars = await clio.get_calendars(user_id=DEMO_USER_ID, writeable=True)
            target_calendar_id = None
            
            if attorney_user_id:
                for cal in calendars:
                    if cal.get("type") == "UserCalendar" and cal.get("permission") == "write":
                        target_calendar_id = cal.get("id")
                        break
            
            await clio.create_calendar_entry(
                user_id=DEMO_USER_ID,
                summary=f"STATUTE OF LIMITATIONS - {accident_details.client_name}",
                start_at=datetime.strptime(sol_date, "%Y-%m-%d"),
                end_at=datetime.strptime(sol_date, "%Y-%m-%d") + timedelta(hours=1),
                matter_id=matter_id, 
                attendee_ids=[attorney_user_id] if attorney_user_id else None,
                all_day=True,
                calendar_owner_id=target_calendar_id
            )
            response.calendar_entry_created = True
        except Exception as e:
            errors.append(f"Calendar entry error: {str(e)}")
            
        # Step 5: Generate Document
        retainer_pdf_content = None
        try:
            doc_result = await clio.create_document_from_template(
                user_id=DEMO_USER_ID,
                template_id=template_id, 
                matter_id=matter_id,
                filename=f"Retainer_Agreement_{accident_details.client_name.replace(' ', '_')}"
            )
            response.document_generated = True
            if doc_result.get("document", {}).get("id"):
                retainer_pdf_content = await clio.download_document(user_id=DEMO_USER_ID, document_id=doc_result["document"]["id"])
        except Exception as e:
            errors.append(f"Document generation error: {str(e)}")

        # Step 6: Email
        if client_email:
            email_result = await email_service.send_client_email(
                client_email=client_email, client_first_name=client_first_name,
                accident_details=accident_details, retainer_pdf_content=retainer_pdf_content
            )
            if email_result["status"] == "sent": response.email_sent = True
            elif email_result["status"] == "preview": response.email_preview = email_result
        
        response.success = len(errors) == 0
        response.errors = errors
        return response
    except Exception as e:
        errors.append(str(e))
        response.errors = errors
        return response

# ==========================================
# Template Generation Endpoint
# ==========================================
@app.post("/api/create-default-template")
async def create_default_template():
    """Generates the default Retainer Agreement and pushes it to Clio as a template."""
    try:
        doc = Document()
        
        # --- NEW: ADD LOGO TO FIRST PAGE HEADER ONLY ---
        section = doc.sections[0]
        # 1. Tell Word this document has a unique first page header
        section.different_first_page_header_footer = True
        
        # 2. Access specifically the FIRST PAGE header
        first_page_header = section.first_page_header
        
        # 3. Get or create the first paragraph in that header
        header_para = first_page_header.paragraphs[0] if first_page_header.paragraphs else first_page_header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 4. Insert the logo if the file exists in the root folder
        logo_path = "logo.png" 
        if os.path.exists(logo_path):
            header_run = header_para.add_run()
            # You can change 2.0 to make it bigger or smaller
            header_run.add_picture(logo_path, width=Inches(2.0))
        
        # --- SMART HELPER FUNCTION ---
        # This automatically finds << Variables >> and makes them bold!
        def add_paragraph_with_vars(document, text):
            p = document.add_paragraph()
            # Split the text by the << >> tags, keeping the tags in the list
            parts = re.split(r'(<<.*?>>)', text)
            for part in parts:
                if part.startswith('<<') and part.endswith('>>'):
                    p.add_run(part).bold = True  # Make the variable bold
                elif part:
                    p.add_run(part)              # Keep normal text normal
            return p
        # -----------------------------

        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("CONTRACT FOR EMPLOYMENT OF ATTORNEYS")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Intro
        add_paragraph_with_vars(doc, "This Retainer Agreement (“Agreement”) is entered into between << Matter.Client.Name >> (“Client”) and << Firm.Name >> (“Attorney”), for the purpose of providing legal representation related to the damages sustained in an incident that occurred on << Matter.CustomField.DateOfAccident >>. By executing this Agreement, Client employs Attorney to investigate, pursue, negotiate, and, if necessary, litigate claims for damages against << Matter.CustomField.DefendantName >> who may be responsible for such damages suffered by Client as a result of << Matter.CustomField.ClientPronounPossessive >> accident.")
        
        # Scope
        add_paragraph_with_vars(doc, "Representation under this Agreement is expressly limited to the matter described herein (“the Claim”) and does not extend to any other legal issues unless separately agreed to in writing by both Client and Attorney. Attorney does not provide tax, accounting, or financial advisory services, and any such issues are outside the scope of this representation. Client is encouraged to consult separate professionals for such matters, as those responsibilities remain << Matter.CustomField.ClientPronounPossessive >> own.")
        
        doc.add_heading("Scope of Representation", level=2)
        add_paragraph_with_vars(doc, "Attorney shall undertake all reasonable and necessary legal efforts to diligently protect and advance Client’s interests in the Claim, extending to both settlement negotiations and litigation proceedings where appropriate. Client agrees to cooperate fully by providing truthful information, timely responses, and all relevant documents or records as requested. Client acknowledges that << Matter.CustomField.ClientPronounPossessive >> cooperation is essential to the effective handling of the Claim.")
        
        # Accident Details
        doc.add_heading("Accident Details & Insurance", level=2)
        add_paragraph_with_vars(doc, "The incident giving rise to this Claim occurred at << Matter.CustomField.AccidentLocation >>. At the time of the accident, Client was operating or occupying a vehicle bearing registration plate number << Matter.CustomField.ClientVehiclePlate >>. The defendant's vehicle was bearing registration plate number << Matter.CustomField.DefendantVehiclePlate >>. The official police report number for this incident is << Matter.CustomField.PoliceReportNumber >>.\n\nThe circumstances surrounding the incident are described as follows: << Matter.CustomField.AccidentDescription >>.\n\nThese circumstances, including the actions of the involved parties and any contributing factors, will be further investigated by Attorney as part of the representation under this Agreement.")
        
        add_paragraph_with_vars(doc, "Attorney is authorized to investigate the liability aspects of the incident, including the collection of police reports, witness statements, and property damage appraisals to determine the full extent of recoverable damages. Client understands that preserving evidence and providing truthful disclosures regarding the events leading to the loss are material obligations under this Agreement. This investigation will serve as the basis for identifying all applicable insurance coverage and responsible parties.")
        
        # --- THE DYNAMIC INJURY CLAUSE ---
        add_paragraph_with_vars(doc, "<< Matter.CustomField.InjuryClause >>")
        
        # Expenses
        doc.add_heading("Litigation Expenses", level=2)
        add_paragraph_with_vars(doc, "Attorney will advance all reasonable costs and expenses necessary for the proper handling of the Claim (“Litigation Expenses”). Such expenses may include, but are not limited to, court filing fees, deposition costs, expert witness fees, medical record retrieval, travel expenses, investigative services, and administrative charges associated with case management.")
        add_paragraph_with_vars(doc, "These Litigation Expenses will be reimbursed to Attorney from Client’s share of the recovery in addition to the contingency fee. Client understands that these expenses are separate from medical bills, liens, or other financial obligations for which << Matter.CustomField.ClientPronounSubject >> may remain personally responsible.")
        
        # Liens
        doc.add_heading("Liens, Subrogation, and Other Obligations", level=2)
        add_paragraph_with_vars(doc, "Client understands that certain parties, such as healthcare providers, insurers, or government agencies (including Medicare or Medicaid), may have a legal right to reimbursement for payments made on Client’s behalf. These are commonly referred to as liens or subrogation claims, and may affect the final amount received by Client from << Matter.CustomField.ClientPronounPossessive >> settlement or judgment. Client hereby authorizes Attorney to negotiate, settle, and satisfy such claims from the proceeds of any recovery. Attorney may engage specialized lien resolution services or other professionals to assist in this process, and the cost of such services shall be treated as a Litigation Expense.")
        
        # SOL
        doc.add_heading("Statute of Limitations", level=2)
        add_paragraph_with_vars(doc, "Attorney will monitor and calculate the deadline for filing the Claim in accordance with applicable law. Based on current information, the statute of limitations for this matter is << Matter.CustomField.StatuteOfLimitations >>. Client acknowledges the importance of timely cooperation in providing documents, records, and information necessary for Attorney to meet all legal deadlines.")
        
        # Termination
        doc.add_heading("Termination of Representation", level=2)
        add_paragraph_with_vars(doc, "Either party may terminate this Agreement upon reasonable written notice. If Client terminates this Agreement after substantial work has been performed, Attorney may assert a claim for attorney’s fees based on the reasonable value of services rendered, payable from any eventual recovery. Client agrees that << Matter.CustomField.ClientPronounPossessive >> obligation to compensate Attorney in such cases shall be limited to the reasonable value of the services rendered up to the point of termination.")
        
        # Signatures
        doc.add_paragraph("\nACCEPTED BY:\n")
        doc.add_paragraph("CLIENT ___________________________                                        Date: _____________________")
        add_paragraph_with_vars(doc, "<< Matter.Client.Name >>\n")
        add_paragraph_with_vars(doc, "<< Firm.Name >> Attorney _________________________             Date: _____________________")
        add_paragraph_with_vars(doc, "<< Matter.ResponsibleAttorney >>")

        # Save to memory and Base64 encode
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        file_base64 = base64.b64encode(file_stream.read()).decode("utf-8")
        
        # Upload to Clio
        result = await clio.create_document_template(
            user_id=DEMO_USER_ID,
            filename="Auto_Generated_Retainer_Agreement.docx",
            file_base64=file_base64
        )
        
        return {"success": True, "data": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@app.get("/health")
def health_check():
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}

@app.post("/api/email-preview")
async def preview_email(data: VerifiedData):
    try:
        accident_details = AccidentDetails(
            date_of_accident=data.date_of_accident,
            accident_location=data.accident_location,
            defendant_name=data.defendant_name,
            client_name=data.client_name,
            client_vehicle_plate=data.client_vehicle_plate,
            defendant_vehicle_plate=data.defendant_vehicle_plate,
            number_injured=data.number_injured,
            accident_description=data.accident_description,
            client_gender=data.client_gender,
            police_report_number=data.police_report_number
        )
        
        subject, html_body = email_service.generate_client_email_content(
            accident_details=accident_details,
            client_first_name=data.client_name.split()[0],
            client_email="preview@example.com"
        )
        
        scheduling_link, link_type = email_service.get_seasonal_scheduling_link()
        
        return {
            "success": True,
            "subject": subject,
            "html_body": html_body,
            "scheduling_link": scheduling_link,
            "scheduling_type": link_type
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/ui", response_class=HTMLResponse)
async def verification_ui():
    """Serve the verification UI for reviewing extracted data."""
    return """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Richards & Law - Police Report Processor</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        .spinner { animation: spin 1s linear infinite; }
        @keyframes spin { to { transform: rotate(360deg); } }
    </style>
</head>
<body class="bg-gray-100 min-h-screen">
    <div class="container mx-auto px-4 py-8">
        <div class="bg-white rounded-lg shadow-md p-6 mb-6">
            <h1 class="text-3xl font-bold text-blue-900">Richards & Law</h1>
            <p class="text-gray-600">Police Report Automation System</p>
        </div>

        <div id="auth-status" class="bg-white rounded-lg shadow-md p-4 mb-6">
            <div class="flex items-center justify-between">
                <span>Clio Connection:</span>
                <span id="auth-indicator" class="px-3 py-1 rounded-full text-sm">Checking...</span>
            </div>
        </div>

        <div class="bg-white rounded-lg shadow-md p-6 mb-6">
            <h2 class="text-xl font-semibold mb-4">Step 1: Upload Police Report</h2>
            <div class="border-2 border-dashed border-gray-300 rounded-lg p-8 text-center" id="drop-zone">
                <input type="file" id="pdf-input" accept=".pdf" class="hidden">
                <label for="pdf-input" class="cursor-pointer">
                    <div class="text-gray-500">
                        <svg class="mx-auto h-12 w-12 text-gray-400" stroke="currentColor" fill="none" viewBox="0 0 48 48">
                            <path d="M28 8H12a4 4 0 00-4 4v20m32-12v8m0 0v8a4 4 0 01-4 4H12a4 4 0 01-4-4v-4m32-4l-3.172-3.172a4 4 0 00-5.656 0L28 28M8 32l9.172-9.172a4 4 0 015.656 0L28 28m0 0l4 4m4-24h8m-4-4v8m-12 4h.02" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" />
                        </svg>
                        <p class="mt-2">Click to upload or drag and drop</p>
                        <p class="text-sm text-gray-400">PDF files only</p>
                    </div>
                </label>
            </div>
            <div id="file-name" class="mt-2 text-sm text-gray-600"></div>
            <button id="extract-btn" class="mt-4 bg-blue-600 text-white px-6 py-2 rounded-md hover:bg-blue-700 disabled:opacity-50" disabled>
                Extract Data with AI
            </button>
            <div id="extract-loading" class="hidden mt-4">
                <div class="flex items-center">
                    <svg class="spinner h-5 w-5 text-blue-600 mr-2" xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24">
                        <circle class="opacity-25" cx="12" cy="12" r="10" stroke="currentColor" stroke-width="4"></circle>
                        <path class="opacity-75" fill="currentColor" d="M4 12a8 8 0 018-8V0C5.373 0 0 5.373 0 12h4zm2 5.291A7.962 7.962 0 014 12H0c0 3.042 1.135 5.824 3 7.938l3-2.647z"></path>
                    </svg>
                    <span>Extracting data with Gemini AI...</span>
                </div>
            </div>
        </div>

        <div id="data-section" class="hidden">
            <div class="bg-white rounded-lg shadow-md p-6 mb-6">
                <h2 class="text-xl font-semibold mb-4">Step 2: Review & Verify Extracted Data</h2>
                <div class="grid grid-cols-1 md:grid-cols-2 gap-4">
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Client Name</label>
                        <input type="text" id="client_name" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                    </div>
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Client Gender</label>
                        <select id="client_gender" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                            <option value="male">Male</option>
                            <option value="female">Female</option>
                        </select>
                    </div>
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Defendant Name</label>
                        <input type="text" id="defendant_name" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                    </div>
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Date of Accident</label>
                        <input type="date" id="date_of_accident" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                    </div>
                    <div class="md:col-span-2">
                        <label class="block text-sm font-medium text-gray-700">Accident Location</label>
                        <input type="text" id="accident_location" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                    </div>
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Client Vehicle Plate</label>
                        <input type="text" id="client_vehicle_plate" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                    </div>
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Defendant Vehicle Plate</label>
                        <input type="text" id="defendant_vehicle_plate" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                    </div>
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Number Injured</label>
                        <input type="number" id="number_injured" min="0" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                    </div>
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Police Report Number</label>
                        <input type="text" id="police_report_number" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                    </div>
                    <div class="md:col-span-2">
                        <label class="block text-sm font-medium text-gray-700">Accident Description</label>
                        <textarea id="accident_description" rows="3" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border"></textarea>
                    </div>
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Statute of Limitations Date (Auto-calculated)</label>
                        <input type="text" id="statute_of_limitations_date" readonly class="mt-1 block w-full rounded-md bg-gray-100 border-gray-300 shadow-sm p-2 border">
                    </div>
                </div>
            </div>

            <div class="bg-white rounded-lg shadow-md p-6 mb-6">
                <h2 class="text-xl font-semibold mb-4">Step 3: Select Clio Matter & Template</h2>
                <div class="grid grid-cols-1 md:grid-cols-2 gap-4">
                    <div>
                        <label class="block text-sm font-medium text-gray-700">Matter</label>
                        <select id="matter_id" class="mt-1 block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                            <option value="">Loading matters...</option>
                        </select>
                    </div>
                    <div>
                        <div class="flex justify-between items-end mb-1">
                            <label class="block text-sm font-medium text-gray-700">Retainer Agreement Template</label>
                            <button id="create-template-btn" class="text-xs bg-indigo-100 text-indigo-700 px-2 py-1 rounded hover:bg-indigo-200 transition-colors">
                                + Add Default Template
                            </button>
                        </div>
                        <select id="template_id" class="block w-full rounded-md border-gray-300 shadow-sm focus:border-blue-500 focus:ring-blue-500 p-2 border">
                            <option value="">Loading templates...</option>
                        </select>
                    </div>
                </div>
            </div>

            <div class="bg-white rounded-lg shadow-md p-6">
                <h2 class="text-xl font-semibold mb-4">Step 4: Process</h2>
                <div class="flex gap-4">
                    <button id="preview-email-btn" class="bg-gray-600 text-white px-6 py-2 rounded-md hover:bg-gray-700">
                        Preview Email
                    </button>
                    <button id="process-btn" class="bg-green-600 text-white px-6 py-2 rounded-md hover:bg-green-700">
                        Process & Send to Clio
                    </button>
                </div>
                <div id="process-loading" class="hidden mt-4">
                    <div class="flex items-center">
                        <svg class="spinner h-5 w-5 text-green-600 mr-2" xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24">
                            <circle class="opacity-25" cx="12" cy="12" r="10" stroke="currentColor" stroke-width="4"></circle>
                            <path class="opacity-75" fill="currentColor" d="M4 12a8 8 0 018-8V0C5.373 0 0 5.373 0 12h4zm2 5.291A7.962 7.962 0 014 12H0c0 3.042 1.135 5.824 3 7.938l3-2.647z"></path>
                        </svg>
                        <span>Processing workflow...</span>
                    </div>
                </div>
            </div>
        </div>

        <div id="results-section" class="hidden mt-6">
            <div class="bg-white rounded-lg shadow-md p-6">
                <h2 class="text-xl font-semibold mb-4">Results</h2>
                <div id="results-content"></div>
            </div>
        </div>

        <div id="email-modal" class="hidden fixed inset-0 bg-black bg-opacity-50 flex items-center justify-center p-4 z-50">
            <div class="bg-white rounded-lg max-w-2xl w-full max-h-[80vh] overflow-auto">
                <div class="p-4 border-b flex justify-between items-center">
                    <h3 class="text-lg font-semibold">Email Preview</h3>
                    <button id="close-modal" class="text-gray-500 hover:text-gray-700">&times;</button>
                </div>
                <div id="email-preview-content" class="p-4"></div>
            </div>
        </div>
    </div>

    <script>
        let extractedData = null;

        // Check auth status
        async function checkAuth() {
            try {
                const resp = await fetch('/oauth/status');
                const data = await resp.json();
                const indicator = document.getElementById('auth-indicator');
                if (data.authenticated) {
                    indicator.textContent = 'Connected';
                    indicator.className = 'px-3 py-1 rounded-full text-sm bg-green-100 text-green-800';
                    loadMattersAndTemplates();
                } else {
                    indicator.innerHTML = '<a href="/oauth/login" class="text-blue-600 underline">Click to connect</a>';
                    indicator.className = 'px-3 py-1 rounded-full text-sm bg-yellow-100 text-yellow-800';
                }
            } catch (e) {
                document.getElementById('auth-indicator').textContent = 'Error';
            }
        }

        // Load matters and templates
        async function loadMattersAndTemplates() {
            try {
                const [mattersResp, templatesResp] = await Promise.all([
                    fetch('/api/matters'),
                    fetch('/api/document-templates')
                ]);
                
                const matters = await mattersResp.json();
                const templates = await templatesResp.json();
                
                const matterSelect = document.getElementById('matter_id');
                matterSelect.innerHTML = '<option value="">Select a matter...</option>';
                matters.data?.forEach(m => {
                    matterSelect.innerHTML += `<option value="${m.id}">${m.display_number} - ${m.description || 'No description'}</option>`;
                });
                
                const templateSelect = document.getElementById('template_id');
                templateSelect.innerHTML = '<option value="">Select a template...</option>';
                templates.data?.forEach(t => {
                    templateSelect.innerHTML += `<option value="${t.id}">${t.name}</option>`;
                });
            } catch (e) {
                console.error('Error loading data:', e);
            }
        }

        // File input handling
        const fileInput = document.getElementById('pdf-input');
        const extractBtn = document.getElementById('extract-btn');
        const fileName = document.getElementById('file-name');

        fileInput.addEventListener('change', (e) => {
            if (e.target.files.length > 0) {
                fileName.textContent = `Selected: ${e.target.files[0].name}`;
                extractBtn.disabled = false;
            }
        });

        // Extract button
        extractBtn.addEventListener('click', async () => {
            const file = fileInput.files[0];
            if (!file) return;

            document.getElementById('extract-loading').classList.remove('hidden');
            extractBtn.disabled = true;

            const formData = new FormData();
            formData.append('file', file);

            try {
                const resp = await fetch('/api/extract-pdf', {
                    method: 'POST',
                    body: formData
                });
                const data = await resp.json();
                
                if (data.success) {
                    extractedData = data.data;
                    populateForm(data.data);
                    document.getElementById('data-section').classList.remove('hidden');
                } else {
                    alert('Extraction failed: ' + data.error);
                }
            } catch (e) {
                alert('Error: ' + e.message);
            } finally {
                document.getElementById('extract-loading').classList.add('hidden');
                extractBtn.disabled = false;
            }
        });

        // Populate form with extracted data
        function populateForm(data) {
            document.getElementById('client_name').value = data.client_name || '';
            document.getElementById('client_gender').value = data.client_gender || 'male';
            document.getElementById('defendant_name').value = data.defendant_name || '';
            document.getElementById('date_of_accident').value = data.date_of_accident || '';
            document.getElementById('accident_location').value = data.accident_location || '';
            document.getElementById('client_vehicle_plate').value = data.client_vehicle_plate || '';
            document.getElementById('defendant_vehicle_plate').value = data.defendant_vehicle_plate || '';
            document.getElementById('number_injured').value = data.number_injured || 0;
            document.getElementById('police_report_number').value = data.police_report_number || '';
            document.getElementById('accident_description').value = data.accident_description || '';
            document.getElementById('statute_of_limitations_date').value = data.statute_of_limitations_date || '';
        }

        // Update SOL date when accident date changes
        document.getElementById('date_of_accident').addEventListener('change', (e) => {
            if (e.target.value) {
                const date = new Date(e.target.value);
                date.setFullYear(date.getFullYear() + 8);
                document.getElementById('statute_of_limitations_date').value = date.toISOString().split('T')[0];
            }
        });

        // Get form data
        function getFormData() {
            return {
                matter_id: parseInt(document.getElementById('matter_id').value),
                client_name: document.getElementById('client_name').value,
                client_gender: document.getElementById('client_gender').value,
                defendant_name: document.getElementById('defendant_name').value,
                date_of_accident: document.getElementById('date_of_accident').value,
                accident_location: document.getElementById('accident_location').value,
                client_vehicle_plate: document.getElementById('client_vehicle_plate').value,
                defendant_vehicle_plate: document.getElementById('defendant_vehicle_plate').value || null,
                number_injured: parseInt(document.getElementById('number_injured').value) || 0,
                police_report_number: document.getElementById('police_report_number').value || null,
                accident_description: document.getElementById('accident_description').value
            };
        }

        // Preview email
        document.getElementById('preview-email-btn').addEventListener('click', async () => {
            const data = getFormData();
            
            try {
                const resp = await fetch('/api/email-preview', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify(data)
                });
                const result = await resp.json();
                
                if (result.success) {
                    document.getElementById('email-preview-content').innerHTML = `
                        <p class="font-semibold mb-2">Subject: ${result.subject}</p>
                        <p class="text-sm text-gray-600 mb-4">Scheduling Link: ${result.scheduling_link} (${result.scheduling_type})</p>
                        <div class="border rounded p-4">${result.html_body}</div>
                    `;
                    document.getElementById('email-modal').classList.remove('hidden');
                }
            } catch (e) {
                alert('Error: ' + e.message);
            }
        });

        // Create Template Button Logic
        document.getElementById('create-template-btn').addEventListener('click', async () => {
            const btn = document.getElementById('create-template-btn');
            btn.disabled = true;
            btn.textContent = "Creating...";
            
            try {
                const resp = await fetch('/api/create-default-template', { method: 'POST' });
                const data = await resp.json();
                
                if (data.success) {
                    alert('Default template successfully created in Clio!');
                    loadMattersAndTemplates(); // Refresh the dropdowns
                } else {
                    alert('Failed to create template: ' + data.error);
                }
            } catch (e) {
                alert('Error: ' + e.message);
            } finally {
                btn.disabled = false;
                btn.textContent = "+ Add Default Template";
            }
        });

        // Close modal
        document.getElementById('close-modal').addEventListener('click', () => {
            document.getElementById('email-modal').classList.add('hidden');
        });

        // Process workflow
        document.getElementById('process-btn').addEventListener('click', async () => {
            const matterId = document.getElementById('matter_id').value;
            const templateId = document.getElementById('template_id').value;
            const file = fileInput.files[0];

            if (!matterId || !templateId) {
                alert('Please select a Matter and Template');
                return;
            }

            document.getElementById('process-loading').classList.remove('hidden');

            const formData = new FormData();
            formData.append('matter_id', matterId);
            formData.append('template_id', templateId);
            formData.append('file', file);

            try {
                const resp = await fetch('/api/workflow', {
                    method: 'POST',
                    body: formData
                });
                const result = await resp.json();
                
                showResults(result);
            } catch (e) {
                alert('Error: ' + e.message);
            } finally {
                document.getElementById('process-loading').classList.add('hidden');
            }
        });

        // Show results
        function showResults(result) {
            const content = document.getElementById('results-content');
            const checkmark = '✓';
            const cross = '✗';
            
            content.innerHTML = `
                <div class="space-y-2">
                    <p class="${result.matter_updated ? 'text-green-600' : 'text-red-600'}">
                        ${result.matter_updated ? checkmark : cross} Matter Updated
                    </p>
                    <p class="${result.custom_fields_set ? 'text-green-600' : 'text-red-600'}">
                        ${result.custom_fields_set ? checkmark : cross} Custom Fields Set
                    </p>
                    <p class="${result.calendar_entry_created ? 'text-green-600' : 'text-red-600'}">
                        ${result.calendar_entry_created ? checkmark : cross} Calendar Entry Created
                    </p>
                    <p class="${result.document_generated ? 'text-green-600' : 'text-red-600'}">
                        ${result.document_generated ? checkmark : cross} Retainer Agreement Generated
                    </p>
                    <p class="${result.email_sent ? 'text-green-600' : 'text-yellow-600'}">
                        ${result.email_sent ? checkmark : '⚠'} Email ${result.email_sent ? 'Sent' : 'Preview Only (SMTP not configured)'}
                    </p>
                </div>
                ${result.errors.length > 0 ? `
                    <div class="mt-4 p-4 bg-red-50 rounded">
                        <p class="font-semibold text-red-800">Errors:</p>
                        <ul class="list-disc list-inside text-red-600">
                            ${result.errors.map(e => `<li>${e}</li>`).join('')}
                        </ul>
                    </div>
                ` : ''}
                ${result.email_preview ? `
                    <div class="mt-4 p-4 bg-blue-50 rounded">
                        <p class="font-semibold text-blue-800">Email Preview:</p>
                        <p>To: ${result.email_preview.to}</p>
                        <p>Subject: ${result.email_preview.subject}</p>
                    </div>
                ` : ''}
            `;
            
            document.getElementById('results-section').classList.remove('hidden');
        }

        // Initialize
        checkAuth();
    </script>
</body>
</html>
"""